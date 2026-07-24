#!/usr/bin/env python3
"""Экспорт DOCX-документов в Markdown.

Скрипт конвертирует DOCX (в первую очередь — результаты OCR журналов/книг) в
Markdown, сохраняя структуру:

* заголовки (стили ``Заголовок №N`` / ``Heading N``) переводятся в ``#``-уровни;
* абзацы основного текста, оглавление, подписи к рисункам/таблицам, сноски;
* полужирный/курсив на уровне отдельных runs;
* таблицы → markdown-таблицы;
* колонтитулы (верхние/нижние) — они в таких файлах лежат в текстовых блоках
  (text box внутри ``w:drawing``), поэтому обычный ``paragraph.text`` их не видит.
  Мы вытаскиваем текст напрямую из XML колонтитула и печатаем его как
  HTML-комментарий на границе каждой секции (≈ страницы).

Результат кладётся рядом с исходным файлом (``foo.docx`` → ``foo.md``), либо в
каталог, указанный через ``--out``.

--------------------------------------------------------------------------------
КАК ЭТО РАБОТАЕТ (обзор пайплайна)
--------------------------------------------------------------------------------
DOCX — это ZIP с XML-частями (``word/document.xml``, ``word/header*.xml`` и т. д.).
python-docx даёт удобный объектный доступ к телу документа, но НЕ ко всему: текст,
лежащий в графических «текстовых блоках» (text box), он не показывает. В OCR-выгрузках
именно так хранятся колонтитулы, поэтому часть работы приходится делать «руками» по
XML через lxml.

``convert_docx_to_md`` идёт по СЕКЦИЯМ документа (в этих файлах одна секция ≈ одна
страница) и для каждой:
  1. печатает верхний колонтитул (если есть) как ``<!-- колонтитул сверху … -->``;
  2. проходит блоки секции (абзацы и таблицы) по порядку через
     ``section.iter_inner_content()`` и конвертирует каждый:
       - заголовки склеиваются (см. буфер ``pending_*``) — многострочный OCR-заголовок
         это несколько абзацев одного стиля подряд, но логически ОДИН заголовок;
       - обычные абзацы → строка Markdown (с учётом стиля: список, подпись, сноска);
       - таблицы → markdown-таблицы;
  3. печатает нижний колонтитул (если есть).
Блоки соединяются пустой строкой (двойной ``\\n``), как того требует Markdown.

Тонкие места, ради которых написан нетривиальный код (подробности — в комментариях
у соответствующих функций):
  * колонтитулы в text box               → ``_extract_header_footer_text`` + XML;
  * дублирование ``mc:Choice``/``mc:Fallback`` → ``_iter_texts`` пропускает Fallback;
  * нестабильный ``id()`` у lxml          → проверка предков через ``iterancestors``;
  * перенос строки ``w:br`` внутри абзаца → ``_runs_to_md`` схлопывает в пробел;
  * многострочный заголовок               → буфер ``pending_*`` в основном цикле.
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# Пространства имён WordprocessingML.
# w  — основной словарь Word (абзацы, runs, текст ``w:t``, колонтитулы ``w:hdr``…).
# mc — Markup Compatibility: механизм «выбор варианта разметки» (AlternateContent),
#      из-за которого один и тот же текст в XML встречается дважды (см. _iter_texts).
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def qn(tag: str) -> str:
    """Возвращает полное («qualified») имя тега с раскрытым неймспейсом.

    lxml хранит теги в форме ``{полный-URI-неймспейса}локальное-имя``. Писать это
    руками неудобно, поэтому передаём короткое ``w:t`` / ``mc:Fallback`` и получаем
    строку вида ``{http://…/main}t``, пригодную для ``element.iter(...)``.
    """

    prefix, local = tag.split(":", 1)
    ns = {"w": W_NS, "mc": MC_NS}[prefix]
    return f"{{{ns}}}{local}"


def _iter_texts(element) -> list[str]:
    """Собирает текст из всех ``w:t`` внутри XML-элемента (в порядке документа).

    Особенность: конструкция ``mc:AlternateContent`` хранит содержимое ДВАЖДЫ — в
    ``mc:Choice`` (современный вариант, напр. DrawingML text box) и в ``mc:Fallback``
    (запасной вариант для старых версий Word, обычно VML с тем же текстом). Если брать
    все ``w:t`` подряд, каждая строка колонтитула удвоится. Поэтому пропускаем всё,
    что лежит внутри ``mc:Fallback``, оставляя только Choice/основной текст.
    """

    fallback_tag = qn("mc:Fallback")

    def _in_fallback(node) -> bool:
        # ВАЖНО про lxml: у элементов НЕТ стабильного id() — на каждый доступ к узлу
        # библиотека создаёт новый Python-proxy, поэтому сравнивать id() бесполезно
        # (сет id-ов из одного обхода не совпадёт с id-ами из другого). Надёжный
        # способ понять «этот w:t внутри Fallback?» — пройтись по его предкам.
        for anc in node.iterancestors():
            if anc.tag == fallback_tag:
                return True
        return False

    texts: list[str] = []
    for t in element.iter(qn("w:t")):
        if _in_fallback(t):
            continue
        texts.append(t.text or "")
    return texts


def _extract_header_footer_text(hf) -> str:
    """Текст колонтитула (верхнего или нижнего) одной строкой.

    ``hf`` — объект ``_Header`` / ``_Footer`` из python-docx. Штатный
    ``hf.paragraphs[..].text`` тут почти всегда пуст: в OCR-DOCX бегущий колонтитул
    (автор на чётной странице, название статьи на нечётной) нарисован через text box
    внутри ``w:drawing``, а не обычными абзацами. Поэтому берём КОРНЕВОЙ XML части
    колонтитула (``hf.part.element`` → ``w:hdr`` / ``w:ftr``) и вытаскиваем текст сами.
    """

    if hf is None:
        return ""
    try:
        element = hf.part.element  # корневой w:hdr / w:ftr
    except Exception:
        # Часть может быть не привязана (пустой/наследуемый колонтитул) — не падаем.
        return ""
    pieces = [p.strip() for p in _iter_texts(element) if p and p.strip()]
    # Схлопываем повторяющиеся ПОДРЯД куски: даже после отсечения Fallback один и тот
    # же фрагмент иногда лежит в нескольких runs подряд (Word дробит текст на runs по
    # смене оформления). Соседние дубликаты убираем, порядок сохраняем.
    deduped: list[str] = []
    for piece in pieces:
        if deduped and deduped[-1] == piece:
            continue
        deduped.append(piece)
    return " ".join(deduped).strip()


def _runs_to_md(paragraph: Paragraph) -> str:
    """Собирает текст абзаца с разметкой полужирного/курсива по отдельным runs.

    Run — это кусок текста абзаца с единым оформлением. Мы идём по runs и для каждого,
    если он жирный/курсивный, оборачиваем в ``**…**`` / ``*…*``. Маркер ставим ВОКРУГ
    обрезанного (без крайних пробелов) текста, а сами пробелы (``lead``/``trail``)
    оставляем снаружи — иначе Markdown-выделение с пробелом у границы «не сработает»
    (``** текст **`` не считается жирным).
    """

    parts: list[str] = []
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        stripped = text.strip()
        if not stripped:
            # Run из одних пробелов — сохраняем как есть (это межсловный пробел).
            parts.append(text)
            continue
        # Отделяем ведущие/замыкающие пробелы, чтобы вынести их за пределы ** / *.
        lead = text[: len(text) - len(text.lstrip())]
        trail = text[len(text.rstrip()) :]
        marker = ""
        if run.bold:
            marker += "**"
        if run.italic:
            marker += "*"
        if marker:
            # marker[::-1] — та же комбинация в обратном порядке для закрытия:
            # открытие ``**​*`` закрывается ``*​**`` (симметрично вложено).
            parts.append(f"{lead}{marker}{stripped}{marker[::-1]}{trail}")
        else:
            parts.append(text)
    result = "".join(parts).strip()
    if not result:
        # В runs текста нет — вероятно, он в text box внутри самого абзаца
        # (напр. врезка). Пытаемся достать его напрямую из XML абзаца (``._p``).
        result = " ".join(t.strip() for t in _iter_texts(paragraph._p) if t.strip()).strip()
    # Схлопываем внутренние переносы строк (``w:br`` приходит как ``\n``) и лишние
    # пробелы: в OCR один логический абзац/заголовок часто разбит переносами по ширине
    # страницы, а в Markdown это должна быть одна логическая строка.
    result = re.sub(r"[ \t\r\n]+", " ", result).strip()
    return result


def _heading_level(style_name: str) -> int | None:
    """Уровень заголовка по имени стиля абзаца, иначе ``None``.

    Понимает и русские OCR-стили (``Заголовок №3``, ``Заголовок 3``), и англ.
    (``Heading 3``). Номер стиля → уровень ``#`` (ограничен диапазоном 1..6). Если
    стиль заголовочный, но без номера — считаем уровнем 2 (разумный дефолт).
    """

    name = (style_name or "").strip().lower()
    for prefix in ("заголовок №", "заголовок ", "heading "):
        if name.startswith(prefix):
            tail = name[len(prefix) :].strip()
            digits = "".join(ch for ch in tail if ch.isdigit())
            if digits:
                return max(1, min(6, int(digits)))
            return 2
    return None


def _paragraph_to_md(paragraph: Paragraph) -> str:
    """Преобразует один НЕ-заголовочный абзац в строку Markdown (или пустую строку).

    Заголовки здесь НЕ обрабатываются — ими занимается основной цикл (склейка
    многострочных заголовков), а сюда попадают только обычные абзацы. Стиль абзаца
    определяет обёртку: оглавление → пункт списка, подписи → курсив, сноска → цитата.
    """

    text = _runs_to_md(paragraph)
    if not text:
        return ""

    style_name = paragraph.style.name if paragraph.style else ""
    lower = style_name.lower()

    # На случай прямого вызова: заголовок всё же оформляем как ``#``-строку.
    level = _heading_level(style_name)
    if level is not None:
        return f"{'#' * level} {text}"

    if "оглавление" in lower:
        return f"- {text}"
    if "подпись к картинке" in lower:
        return f"*{text}*"
    if "подпись к таблице" in lower:
        return f"*{text}*"
    if "сноска" in lower:
        return f"> {text}"

    return text


def _table_to_md(table: Table) -> str:
    """Преобразует таблицу в markdown-таблицу (первая строка — шапка).

    Особые случаи: внутри ячейки может быть несколько абзацев — склеиваем через
    ``<br>`` (перенос внутри ячейки Markdown); символ ``|`` экранируем, чтобы не
    сломать разметку столбцов; строки с разным числом ячеек дополняем пустыми до
    общей ширины (иначе таблица «съедет»).
    """

    rows: list[list[str]] = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = "<br>".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cells.append(cell_text.replace("|", "\\|").strip())
        rows.append(cells)

    if not rows:
        return ""

    # Выравниваем все строки по максимальному числу столбцов.
    ncols = max(len(r) for r in rows)
    rows = [r + [""] * (ncols - len(r)) for r in rows]

    header = rows[0]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(["---"] * ncols) + " |"]
    for r in rows[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def convert_docx_to_md(docx_path: Path, with_colontituls: bool = True) -> str:
    """Конвертирует DOCX в текст Markdown и возвращает его строкой.

    Идём по секциям (≈ страницам). Внутри секции обходим блоки по порядку и
    аккуратно склеиваем многострочные заголовки (см. ниже), а колонтитулы печатаем
    до и после содержимого секции как HTML-комментарии.
    """

    document = Document(str(docx_path))
    out: list[str] = []

    for sec_idx, section in enumerate(document.sections):
        if with_colontituls:
            # У секции может быть до трёх вариантов колонтитула (обычный / первая
            # страница / чётные страницы). Берём первый непустой.
            header_text = _extract_header_footer_text(section.header)
            first_header = _extract_header_footer_text(section.first_page_header)
            even_header = _extract_header_footer_text(section.even_page_header)
            top = header_text or first_header or even_header
            if top:
                out.append(f"<!-- колонтитул сверху [секция {sec_idx + 1}]: {top} -->")

        # --- Склейка многострочного заголовка --------------------------------
        # OCR разбивает длинный заголовок на несколько абзацев ОДНОГО стиля,
        # идущих подряд. Логически это один заголовок. Копим их части в буфере
        # pending_parts, пока подряд идут заголовки того же уровня; на первом же
        # не-заголовке (или заголовке другого уровня, или таблице, или конце
        # секции) «сбрасываем» буфер одной ``#``-строкой через flush_heading().
        pending_level: int | None = None
        pending_parts: list[str] = []

        def flush_heading() -> None:
            nonlocal pending_level, pending_parts
            if pending_level is not None and pending_parts:
                text = " ".join(pending_parts)
                # После склейки двух выделенных кусков получается ``**A** **B**``.
                # Сливаем соседние одинаковые выделения в одно: ``**A B**`` (и то же
                # для одиночного курсива ``*A* *B*`` → ``*A B*``).
                text = re.sub(r"\*\*\s+\*\*", " ", text)
                text = re.sub(r"(?<!\*)\*\s+\*(?!\*)", " ", text)
                out.append(f"{'#' * pending_level} {text}")
            pending_level = None
            pending_parts = []

        for block in section.iter_inner_content():
            if isinstance(block, Paragraph):
                level = _heading_level(block.style.name if block.style else "")
                if level is not None:
                    inner = _runs_to_md(block)
                    if not inner:
                        continue
                    if pending_level == level:
                        # Продолжение того же заголовка — добавляем строку в буфер.
                        pending_parts.append(inner)
                    else:
                        # Начался заголовок другого уровня — закрываем прежний.
                        flush_heading()
                        pending_level = level
                        pending_parts = [inner]
                    continue
                # Обычный абзац: сначала закрываем висящий заголовок, потом печатаем.
                flush_heading()
                md = _paragraph_to_md(block)
                if md:
                    out.append(md)
            elif isinstance(block, Table):
                flush_heading()
                md = _table_to_md(block)
                if md:
                    out.append(md)

        # Конец секции — не даём заголовку «утечь» в следующую секцию.
        flush_heading()

        if with_colontituls:
            footer_text = _extract_header_footer_text(section.footer)
            first_footer = _extract_header_footer_text(section.first_page_footer)
            even_footer = _extract_header_footer_text(section.even_page_footer)
            bottom = footer_text or first_footer or even_footer
            if bottom:
                out.append(f"<!-- колонтитул снизу [секция {sec_idx + 1}]: {bottom} -->")

    # Блоки разделяем пустой строкой (в Markdown это граница абзацев/элементов).
    text = "\n\n".join(out).strip() + "\n"
    return text


def _iter_docx_files(paths: list[Path], recursive: bool) -> list[Path]:
    """Разворачивает список путей (файлы/каталоги) в плоский список .docx-файлов.

    Каталоги обходятся glob-ом (рекурсивно при ``--recursive``). Временные файлы Word
    (``~$имя.docx`` — блокировки открытого документа) пропускаем.
    """

    result: list[Path] = []
    for p in paths:
        if p.is_dir():
            pattern = "**/*.docx" if recursive else "*.docx"
            result.extend(sorted(q for q in p.glob(pattern) if not q.name.startswith("~$")))
        elif p.suffix.lower() == ".docx":
            result.append(p)
        else:
            print(f"Пропускаю (не .docx): {p}", file=sys.stderr)
    return result


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Экспорт DOCX в Markdown.")
    parser.add_argument("paths", nargs="+", type=Path, help="DOCX-файлы или каталоги с ними")
    parser.add_argument(
        "--out", type=Path, default=None, help="Каталог для .md (по умолчанию — рядом с исходным файлом)"
    )
    parser.add_argument("--recursive", action="store_true", help="Искать .docx в подкаталогах")
    parser.add_argument("--no-colontituls", action="store_true", help="Не выгружать колонтитулы (верхние/нижние)")
    args = parser.parse_args(argv)

    files = _iter_docx_files(args.paths, args.recursive)
    if not files:
        print("Не найдено ни одного .docx", file=sys.stderr)
        return 1

    if args.out is not None:
        args.out.mkdir(parents=True, exist_ok=True)

    exit_code = 0
    for docx_path in files:
        try:
            md = convert_docx_to_md(docx_path, with_colontituls=not args.no_colontituls)
        except Exception as exc:  # noqa: BLE001 — хотим продолжить с остальными файлами
            print(f"ОШИБКА при обработке {docx_path}: {exc}", file=sys.stderr)
            exit_code = 1
            continue

        # По умолчанию кладём .md рядом с .docx; с --out — в указанный каталог.
        if args.out is not None:
            md_path = args.out / (docx_path.stem + ".md")
        else:
            md_path = docx_path.with_suffix(".md")
        md_path.write_text(md, encoding="utf-8")
        print(f"{docx_path}  ->  {md_path}  ({len(md)} символов)")

    return exit_code


if __name__ == "__main__":
    raise SystemExit(main())
